import fitz
from typing import List, Tuple
from pathlib import Path
import base64
import hashlib
import os

from models.schemas import ContentBlock
from utils.exceptions import ParseError
from config import DATA_DIR


IMAGES_DIR = DATA_DIR / "images"
IMAGES_DIR.mkdir(parents=True, exist_ok=True)


def _save_image(image_bytes: bytes, doc_hash: str, image_index: int, image_ext: str = "png") -> str:
    # 统一图片格式后缀，去掉多余的点
    if image_ext.startswith('.'):
        image_ext = image_ext[1:]
    image_ext = image_ext.lower() or "png"
    
    # 只保留常见图片格式，未知格式默认png
    allowed_exts = {"png", "jpg", "jpeg", "gif", "bmp", "webp"}
    if image_ext not in allowed_exts:
        image_ext = "png"
    
    image_filename = f"{doc_hash}_img_{image_index}.{image_ext}"
    image_path = IMAGES_DIR / image_filename
    
    with open(image_path, "wb") as f:
        f.write(image_bytes)
    
    return f"/api/images/{image_filename}"


def extract_text_from_pdf(pdf_path: str) -> Tuple[str, int, List[ContentBlock]]:
    try:
        doc = fitz.open(pdf_path)
        total_pages = len(doc)
        blocks = []
        full_text = ""
        
        doc_hash = hashlib.md5(Path(pdf_path).name.encode()).hexdigest()[:8]
        image_index = 0
        
        for page_num in range(total_pages):
            page = doc[page_num]
            
            # 获取页面所有元素（文本块+图片），按y坐标排序
            page_elements = []
            
            # 先获取文本块
            text_blocks = page.get_text("dict")["blocks"]
            for b in text_blocks:
                if b["type"] == 0:  # 文本块
                    text_content = ""
                    for line in b["lines"]:
                        for span in line["spans"]:
                            text_content += span["text"]
                        text_content += "\n"
                    if text_content.strip():
                        page_elements.append({
                            "type": "text",
                            "y": b["bbox"][1],  # 上边界y坐标
                            "content": text_content.strip()
                        })
            
            # 再获取图片块
            image_list = page.get_images(full=True)
            for img_info in image_list:
                xref = img_info[0]
                try:
                    base_image = doc.extract_image(xref)
                    image_bytes = base_image["image"]
                    image_ext = base_image.get("ext", "png")
                    
                    # 获取图片位置
                    img_rects = page.get_image_rects(xref)
                    if img_rects:
                        y_pos = img_rects[0][1]  # 图片上边界y坐标
                    else:
                        y_pos = 99999  # 未知位置放最后
                    
                    if len(image_bytes) > 1000:
                        image_index += 1
                        image_url = _save_image(image_bytes, doc_hash, image_index, image_ext)
                        md_content = f"![图片 {image_index}]({image_url})"
                        
                        page_elements.append({
                            "type": "image",
                            "y": y_pos,
                            "content": md_content,
                            "image_path": image_url,
                            "image_id": f"img_{page_num + 1}_{image_index}"
                        })
                except Exception:
                    pass
            
            # 按y坐标从小到大排序，保证从上到下的顺序
            page_elements.sort(key=lambda x: x["y"])
            
            # 处理页面标题
            full_text += f"\n--- 第 {page_num + 1} 页 ---\n"
            
            # 按顺序添加元素
            for elem in page_elements:
                if elem["type"] == "text":
                    full_text += elem["content"] + "\n\n"
                    blocks.append(ContentBlock(
                        id=f"page_{page_num + 1}_{len(blocks)}",
                        type="page",
                        page=page_num + 1,
                        content=elem["content"],
                    ))
                elif elem["type"] == "image":
                    full_text += elem["content"] + "\n\n"
                    blocks.append(ContentBlock(
                        id=elem["image_id"],
                        type="image",
                        page=page_num + 1,
                        content=elem["content"],
                        image_path=elem["image_path"],
                    ))
        
        doc.close()
        return full_text.strip(), total_pages, blocks
        
    except Exception as e:
        raise ParseError(f"PDF 文本提取失败: {str(e)}")


def extract_text_from_docx(docx_path: str) -> Tuple[str, int, List[ContentBlock]]:
    """提取 DOCX 内容并转换为 Markdown 格式"""
    try:
        from docx import Document
        from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
        
        doc = Document(docx_path)
        blocks = []
        full_text = ""
        paragraph_count = 0
        
        doc_hash = hashlib.md5(Path(docx_path).name.encode()).hexdigest()[:8]
        image_index = 0
        
        # 提取图片关系
        image_rels = {}
        for rel in doc.part.rels.values():
            if "image" in rel.reltype:
                image_rels[rel.rId] = rel.target_part.blob
        
        # 处理段落
        for para in doc.paragraphs:
            text = para.text.strip()
            
            # 识别标题样式
            style_name = para.style.name if para.style else ""
            is_heading = False
            heading_level = 0
            
            if style_name.startswith('Heading'):
                try:
                    heading_level = int(style_name.split()[-1])
                    is_heading = True
                except:
                    pass
            
            # 处理段落中的图片
            para_images = []
            for run in para.runs:
                run_xml = run._element.xml
                if 'drawing' in run_xml or 'blip' in run_xml:
                    import re
                    rIds = re.findall(r'r:embed="([^"]+)"', run_xml)
                    for rId in rIds:
                        if rId in image_rels:
                            image_bytes = image_rels[rId]
                            if len(image_bytes) > 1000:
                                # 识别图片格式
                                if image_bytes.startswith(b'\xff\xd8\xff'):
                                    image_ext = "jpg"
                                elif image_bytes.startswith(b'\x89PNG\r\n\x1a\n'):
                                    image_ext = "png"
                                elif image_bytes.startswith(b'GIF89a') or image_bytes.startswith(b'GIF87a'):
                                    image_ext = "gif"
                                else:
                                    image_ext = "png"
                                
                                image_index += 1
                                image_url = _save_image(image_bytes, doc_hash, image_index, image_ext)
                                
                                blocks.append(ContentBlock(
                                    id=f"img_{image_index}",
                                    type="image",
                                    page=1,
                                    content=f"![图片 {image_index}]({image_url})",
                                    image_path=image_url,
                                ))
                                
                                para_images.append(f"\n![图片 {image_index}]({image_url})\n")
            
            # 转换为 Markdown 格式
            if text:
                paragraph_count += 1
                
                # 标题处理
                if is_heading and heading_level > 0:
                    prefix = "#" * heading_level
                    md_text = f"{prefix} {text}"
                # 列表处理（识别项目符号）
                elif para.style and 'List' in para.style.name:
                    md_text = f"- {text}"
                else:
                    md_text = text
                
                full_text += md_text + "\n\n"
                blocks.append(ContentBlock(
                    id=f"para_{paragraph_count}",
                    type="paragraph",
                    page=1,
                    content=md_text,
                ))
            
            # 添加段落中的图片
            for img_md in para_images:
                full_text += img_md
        
        # 处理表格（转换为 Markdown 表格）
        for table in doc.tables:
            if len(table.rows) == 0:
                continue
            
            # 表格标题
            full_text += "\n"
            
            # 表头
            header_row = table.rows[0]
            headers = [cell.text.strip() for cell in header_row.cells]
            table_md = "| " + " | ".join(headers) + " |\n"
            table_md += "|" + "|".join(["---" for _ in headers]) + "|\n"
            
            # 表格内容
            for row in table.rows[1:]:
                cells = [cell.text.strip() for cell in row.cells]
                table_md += "| " + " | ".join(cells) + " |\n"
            
            full_text += table_md + "\n"
            blocks.append(ContentBlock(
                id=f"table_{len(blocks) + 1}",
                type="table",
                page=1,
                content=table_md.strip(),
            ))
        
        return full_text.strip(), 1, blocks
        
    except ImportError:
        raise ParseError("需要安装 python-docx 库来处理 docx 文件")
    except Exception as e:
        raise ParseError(f"DOCX 文本提取失败: {str(e)}")


def extract_text_from_doc(doc_path: str) -> Tuple[str, int, List[ContentBlock]]:
    """提取 DOC 内容并转换为 Markdown 格式"""
    try:
        import subprocess
        import tempfile
        
        docx_path = Path(doc_path).with_suffix('.docx')
        
        # 尝试使用 LibreOffice 转换为 docx
        try:
            subprocess.run([
                'soffice', '--headless', '--convert-to', 'docx',
                '--outdir', str(Path(doc_path).parent),
                str(doc_path)
            ], check=True, capture_output=True, timeout=60)
            
            if docx_path.exists():
                # 转换成功，使用 docx 解析（已转换为 Markdown）
                result = extract_text_from_docx(str(docx_path))
                try:
                    docx_path.unlink()
                except:
                    pass
                return result
        except (subprocess.TimeoutExpired, subprocess.CalledProcessError):
            pass
        except FileNotFoundError:
            pass
        
        # LibreOffice 不可用，尝试使用 pywin32（Windows）
        try:
            import win32com.client
            word = win32com.client.Dispatch("Word.Application")
            word.Visible = False
            
            doc = word.Documents.Open(str(Path(doc_path).absolute()))
            full_text = doc.Content.Text
            doc.Close(False)
            word.Quit()
            
            # 简单转换为 Markdown 格式
            blocks = []
            paragraphs = full_text.split('\n')
            md_text = ""
            
            for i, para in enumerate(paragraphs):
                if para.strip():
                    # 尝试识别标题（通常标题较短且可能有大写）
                    if len(para.strip()) < 50 and para.strip().startswith('第') or para.strip().endswith('章'):
                        md_para = f"## {para.strip()}"
                    else:
                        md_para = para.strip()
                    
                    md_text += md_para + "\n\n"
                    blocks.append(ContentBlock(
                        id=f"para_{i + 1}",
                        type="paragraph",
                        page=1,
                        content=md_para,
                    ))
            
            return md_text.strip(), 1, blocks
            
        except ImportError:
            raise ParseError("需要安装 pywin32 库或 LibreOffice 来处理 doc 文件")
        
    except Exception as e:
        raise ParseError(f"DOC 文本提取失败: {str(e)}")


def extract_text_from_md(md_path: str) -> Tuple[str, int, List[ContentBlock]]:
    """直接读取 Markdown 文件，无需转换"""
    try:
        with open(md_path, 'r', encoding='utf-8') as f:
            content = f.read()
        
        # 创建 blocks（按段落分割）
        blocks = []
        lines = content.split('\n\n')  # 按双换行分割段落
        for i, para in enumerate(lines):
            if para.strip():
                blocks.append(ContentBlock(
                    id=f"para_{i + 1}",
                    type="paragraph",
                    page=1,
                    content=para.strip(),
                ))
        
        return content, 1, blocks
        
    except Exception as e:
        raise ParseError(f"Markdown 文件读取失败: {str(e)}")


def extract_text_from_document(file_path: str) -> Tuple[str, int, List[ContentBlock]]:
    path = Path(file_path)
    ext = path.suffix.lower()
    
    if ext == '.pdf':
        return extract_text_from_pdf(file_path)
    elif ext == '.docx':
        return extract_text_from_docx(file_path)
    elif ext == '.doc':
        return extract_text_from_doc(file_path)
    elif ext == '.md':
        return extract_text_from_md(file_path)
    else:
        raise ParseError(f"不支持的文件格式: {ext}")